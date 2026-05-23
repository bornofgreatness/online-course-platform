"""Generate 1-page client launch outline DOCX (EN + PT-BR)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_DIR = Path(__file__).resolve().parents[1] / "docs"


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)


def add_title(doc: Document, text: str, subtitle: str) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(subtitle)
    sr.font.size = Pt(9)
    sr.italic = True
    doc.add_paragraph()


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")


def add_check(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"☐  {text}")


def build_en() -> Document:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Going Live — Your Checklist",
        "Online Course Platform · Client-owned accounts · 1-page outline",
    )

    add_heading(doc, "Why you buy these (in your name)")
    add_bullet(
        doc,
        "AWS = rent for the website (24/7) + safe storage for course PDFs/videos.",
    )
    add_bullet(
        doc,
        "Domain + Resend = your web address + emails from you (e.g. noreply@yourdomain.com).",
    )
    add_bullet(doc, "Developer installs the app; you own accounts, bills, and data.")

    add_heading(doc, "What to create & pay for")
    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    hdr = ["#", "Service", "Your action"]
    rows = [
        ("1", "Domain (e.g. .com.br)", "Buy at Registro.br or registrar; save login."),
        ("2", "AWS", "Create account + card; billing alerts on."),
        ("3", "Resend", "Sign up; add domain; paste DNS until Verified."),
        ("4", "Supabase (database)", "New project; save DB password + connection URLs."),
        ("5", "Stripe", "Complete KYC; live API keys when ready."),
        ("6", "DNS for website", "After dev gives IP: A record → server (call if needed)."),
    ]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    add_heading(doc, "Your checklist")
    for item in [
        "Domain purchased: _________________________",
        "AWS account created (your email/card)",
        "Resend domain verified",
        "Supabase project created",
        "Stripe live account ready",
        "Developer access sent securely (not WhatsApp)",
    ]:
        add_check(doc, item)

    add_heading(doc, "Send to developer (secure)")
    add_bullet(doc, "Domain name · AWS keys or IAM user · Resend API key (re_...)")
    add_bullet(doc, "Sender email · Supabase URLs · Stripe pk_live / sk_live · WhatsApp #")

    add_heading(doc, "Rough monthly cost (start)")
    add_bullet(doc, "AWS hosting + files: ~R$ 100–250 · Database: R$ 0–130 · Resend: often free at low volume")
    add_bullet(doc, "Domain: yearly · Stripe: fee per sale only")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Client name: _________________  Date: _________  Signature: _________________")
    r.font.size = Pt(9)
    return doc


def build_pt() -> Document:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Colocar no Ar — Seu Checklist",
        "Plataforma de Cursos · Contas no seu nome · Resumo de 1 página",
    )

    add_heading(doc, "Por que você contrata (no seu nome)")
    add_bullet(
        doc,
        "AWS = aluguel do site (24h) + armazenamento seguro de PDFs/vídeos dos cursos.",
    )
    add_bullet(
        doc,
        "Domínio + Resend = seu endereço na web + e-mails seus (ex.: noreply@seudominio.com.br).",
    )
    add_bullet(doc, "O desenvolvedor instala o sistema; você é dono das contas, faturas e dados.")

    add_heading(doc, "O que criar e pagar")
    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    hdr = ["#", "Serviço", "O que você faz"]
    rows = [
        ("1", "Domínio (ex. .com.br)", "Comprar no Registro.br; guardar login."),
        ("2", "AWS", "Criar conta + cartão; ativar alertas de fatura."),
        ("3", "Resend", "Cadastrar; adicionar domínio; colar DNS até Verificado."),
        ("4", "Supabase (banco de dados)", "Novo projeto; salvar senha e URLs de conexão."),
        ("5", "Stripe", "Completar cadastro; chaves live quando for ao ar."),
        ("6", "DNS do site", "Após IP do servidor: registro A (vide chamada se precisar)."),
    ]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    add_heading(doc, "Seu checklist")
    for item in [
        "Domínio comprado: _________________________",
        "Conta AWS criada (seu e-mail/cartão)",
        "Domínio verificado no Resend",
        "Projeto Supabase criado",
        "Stripe em modo live pronto",
        "Acessos enviados ao dev com segurança (não WhatsApp)",
    ]:
        add_check(doc, item)

    add_heading(doc, "Enviar ao desenvolvedor (com segurança)")
    add_bullet(doc, "Domínio · chaves AWS ou usuário IAM · chave Resend (re_...)")
    add_bullet(doc, "E-mail remetente · URLs Supabase · Stripe pk_live / sk_live · WhatsApp")

    add_heading(doc, "Custo mensal aproximado (início)")
    add_bullet(doc, "AWS site + arquivos: ~R$ 100–250 · Banco: R$ 0–130 · Resend: grátis no começo")
    add_bullet(doc, "Domínio: anual · Stripe: taxa só por venda")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cliente: _________________  Data: _________  Assinatura: _________________")
    r.font.size = Pt(9)
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    en_path = OUT_DIR / "CLIENT_LAUNCH_CHECKLIST_EN.docx"
    pt_path = OUT_DIR / "CLIENT_LAUNCH_CHECKLIST_PT-BR.docx"
    build_en().save(en_path)
    build_pt().save(pt_path)
    print(f"Created: {en_path}")
    print(f"Created: {pt_path}")


if __name__ == "__main__":
    main()
